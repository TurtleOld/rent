import datetime
import locale
import os
import tempfile
from decimal import Decimal

from dateutil import parser
from django.contrib.messages.views import SuccessMessageMixin
from django.urls import reverse_lazy
from django.views.generic import TemplateView, FormView

from collect.custom_mixins import CustomNoPermissionMixin
from collect.rent.forms import UploadFileForm
from docx import Document
from pdf2docx import Converter

from collect.rent.models import Rent, ServiceInfo

TYPE_SERVICE = [
    'ВЗНОС НА КАП. РЕМОНТ',
    'ВОДООТВЕДЕНИЕ ОДН',
    'ГОРЯЧАЯ ВОДА (НОСИТЕЛЬ) ОДН',
    'ГОРЯЧЕЕ В/С (ЭНЕРГИЯ) ОДН',
    'ГОРЯЧЕЕ В/С (НОСИТЕЛЬ) ОДН',
    'СОДЕРЖАНИЕ Ж/Ф',
    'ХОЛОДНОЕ В/С ОДН',
    'ЭЛЕКТРОЭНЕРГИЯ ОДН',
    'ВОДООТВЕДЕНИЕ',
    'ГАЗОСНАБЖЕНИЕ',
    'ГОРЯЧЕЕ  В/С (ЭНЕРГИЯ).',
    'ГОРЯЧЕЕ В/С (НОСИТЕЛЬ)',
    'ОБРАЩЕНИЕ С ТКО',
    'ОТОПЛЕНИЕ',
    'ХОЛОДНОЕ В/С',
    'ДОБРОВОЛЬНОЕ СТРАХОВАНИЕ',
    'ЗАПИРАЮЩЕЕ УСТРОЙСТВО',
]

MONTHS = {
    'Январь': 1,
    'Февраль': 2,
    'Март': 3,
    'Апрель': 4,
    'Май': 5,
    'Июнь': 6,
    'Июль': 7,
    'Август': 8,
    'Сентябрь': 9,
    'Октябрь': 10,
    'Ноябрь': 11,
    'Декабрь': 12,
}


class RentView(CustomNoPermissionMixin, SuccessMessageMixin, TemplateView):
    template_name = 'rent/index.html'
    no_permission_url = reverse_lazy('login')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Привет!'
        return context


class FileFieldFormView(CustomNoPermissionMixin, SuccessMessageMixin, FormView):
    template_name = 'rent/payslips.html'
    form_class = UploadFileForm
    success_url = reverse_lazy('rent:list')

    def post(self, request, *args, **kwargs):
        form_class = self.get_form_class()
        form = self.get_form(form_class)
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)

    def form_valid(self, form):
        file = form.cleaned_data['file']
        with tempfile.NamedTemporaryFile(mode='w+b') as tempf:
            tempf.write(file.read())
            docx_file = convert_pdf_to_docx(tempf)
            format_rent(docx_file)
            return super().form_valid(form)


def convert_pdf_to_docx(file):

    path_input = file
    path_output = 'output.docx'

    cv = Converter(path_input)
    cv.convert(path_output, start=0, end=None, multi_proccessing=True)
    cv.close()
    return path_output


locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')


def format_rent(file):
    document = Document(file)
    personal_account = ''
    date_str = document.paragraphs[1].text.split('ЗА')[1].strip().capitalize()
    month_str, year_str = date_str[:-2].split()
    month = MONTHS[month_str]
    year = int(year_str)
    date = datetime.datetime(year, month, 1).date()
    for paragraph in document.paragraphs:
        if 'Кому:' in paragraph.text:
            personal_account = (
                paragraph.text.split('Кому:')[1].split('Куда:')[0].strip()
            )
    rent_info, _ = Rent.objects.get_or_create(personal_account=personal_account)
    for item in document.tables[3].rows:
        text = [cell.text for cell in item.cells]
        if text[0] in TYPE_SERVICE:
            ServiceInfo.objects.create(
                rent=rent_info,
                date=date,
                type_service=text[0],
                scope_service=Decimal(
                    text[1].replace(',', '.').replace(' ', ''),
                ),
                units=text[2],
                tariff=Decimal(text[3].replace(',', '.').replace(' ', '')),
                accrued_tariff=Decimal(
                    text[4].replace(',', '.').replace(' ', ''),
                ),
                recalculations=(
                    Decimal(
                        text[5].replace(',', '.').replace(' ', ''),
                    )
                    if len(text) > 6
                    else 0
                ),
                total=Decimal(text[-1].replace(',', '.').replace(' ', '')),
            )
